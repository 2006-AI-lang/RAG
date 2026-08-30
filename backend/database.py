"""SQLite 数据库初始化与管理。"""

import sqlite3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional
from config import settings


DB_PATH = Path(settings.DATABASE_URL.replace("sqlite:///", ""))
if not DB_PATH.is_absolute():
    DB_PATH = Path(__file__).parent / DB_PATH


def get_connection() -> sqlite3.Connection:
    """获取 SQLite 连接。"""
    conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    """初始化数据库表。"""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS qa_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question TEXT NOT NULL,
            answer TEXT NOT NULL,
            mode TEXT NOT NULL DEFAULT 'hybrid',
            sources TEXT DEFAULT NULL,
            user_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS llm_models (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            base_url TEXT NOT NULL,
            api_key_encrypted TEXT NOT NULL,
            model_name TEXT NOT NULL,
            is_active INTEGER DEFAULT 0,
            user_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # 迁移：旧库无 user_id 列时补充（NULL 表示全局模型）
    try:
        cursor.execute("ALTER TABLE llm_models ADD COLUMN user_id INTEGER")
    except sqlite3.OperationalError:
        pass  # 列已存在

    # 迁移：移除 name 的全局 UNIQUE，改为 (user_id, name) 唯一，允许不同用户/全局同名
    indexes = cursor.execute("PRAGMA index_list(llm_models)").fetchall()
    if any(idx[1].startswith("sqlite_autoindex_llm_models") for idx in indexes):
        cursor.execute("""
            CREATE TABLE llm_models_mig (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                base_url TEXT NOT NULL,
                api_key_encrypted TEXT NOT NULL,
                model_name TEXT NOT NULL,
                is_active INTEGER DEFAULT 0,
                user_id INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cursor.execute(
            """
            INSERT INTO llm_models_mig (id, name, base_url, api_key_encrypted, model_name, is_active, user_id, created_at)
            SELECT id, name, base_url, api_key_encrypted, model_name, is_active, user_id, created_at FROM llm_models
            """
        )
        cursor.execute("DROP TABLE llm_models")
        cursor.execute("ALTER TABLE llm_models_mig RENAME TO llm_models")

    # (user_id, name) 唯一索引；不同用户或全局与个人可同名
    cursor.execute(
        "CREATE UNIQUE INDEX IF NOT EXISTS ux_llm_models_user_name ON llm_models(user_id, name)"
    )

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS knowledge_entries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entry_id TEXT NOT NULL UNIQUE,
            title TEXT NOT NULL,
            category TEXT NOT NULL DEFAULT '未分类',
            content TEXT NOT NULL,
            source TEXT NOT NULL DEFAULT '',
            url TEXT DEFAULT '',
            tags TEXT DEFAULT '[]',
            import_mode TEXT DEFAULT 'direct',
            user_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS auth_tokens (
            token TEXT PRIMARY KEY,
            user_id INTEGER NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            expires_at TIMESTAMP DEFAULT NULL,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    # 迁移：为旧 auth_tokens 添加 expires_at 列
    try:
        cursor.execute("ALTER TABLE auth_tokens ADD COLUMN expires_at TIMESTAMP DEFAULT NULL")
    except sqlite3.OperationalError:
        pass

    # 迁移：为旧 qa_history 添加 user_id 列
    try:
        cursor.execute("ALTER TABLE qa_history ADD COLUMN user_id INTEGER")
    except sqlite3.OperationalError:
        pass

    # 迁移：为旧 knowledge_entries 添加 user_id 列
    try:
        cursor.execute("ALTER TABLE knowledge_entries ADD COLUMN user_id INTEGER")
    except sqlite3.OperationalError:
        pass

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS chat_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL DEFAULT '新对话',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            sources TEXT DEFAULT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (session_id) REFERENCES chat_sessions(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS unanswered_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question TEXT NOT NULL,
            mode TEXT NOT NULL DEFAULT 'hybrid',
            reason TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS user_llm_config (
            user_id INTEGER PRIMARY KEY,
            base_url TEXT NOT NULL DEFAULT '',
            api_key_encrypted TEXT NOT NULL DEFAULT '',
            model_name TEXT NOT NULL DEFAULT '',
            mode TEXT NOT NULL DEFAULT 'mock',
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS exercise_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            exercise_type TEXT NOT NULL,
            duration INTEGER DEFAULT 0,
            intensity TEXT DEFAULT '中等',
            notes TEXT DEFAULT '',
            record_date TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS knowledge_versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entry_id TEXT NOT NULL,
            version INTEGER NOT NULL DEFAULT 1,
            title TEXT NOT NULL,
            category TEXT NOT NULL DEFAULT '未分类',
            content TEXT NOT NULL,
            source TEXT NOT NULL DEFAULT '',
            url TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS user_retrieval_config (
            user_id INTEGER PRIMARY KEY,
            default_mode TEXT NOT NULL DEFAULT 'hybrid',
            default_top_k INTEGER NOT NULL DEFAULT 5,
            min_vector_score REAL NOT NULL DEFAULT 0.30,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS training_plans (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL DEFAULT '',
            goal TEXT NOT NULL DEFAULT '',
            level TEXT NOT NULL DEFAULT '',
            days_per_week INTEGER DEFAULT 4,
            content TEXT NOT NULL DEFAULT '',
            category TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS category_keywords (
            category TEXT PRIMARY KEY,
            keywords TEXT NOT NULL DEFAULT '[]',
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # 首次初始化时插入默认分类关键词
    cursor.execute("SELECT COUNT(*) FROM category_keywords")
    if cursor.fetchone()[0] == 0:
        defaults = [
            ("力量训练", '["深蹲", "硬拉", "卧推", "推举", "划船", "哑铃", "杠铃", "力量", "RM", "复合动作", "三大项", "孤立动作"]'),
            ("增肌", '["增肌", "肌肉增长", "肌肥大", "增重", "维度", "肌肉量", "肌肉线条"]'),
            ("减脂", '["减脂", "减肥", "瘦身", "燃脂", "热量消耗", "体重下降", "体脂率", "瘦体重"]'),
            ("有氧运动", '["跑步", "游泳", "骑行", "跳绳", "椭圆机", "有氧", "HIIT", "心肺", "耐力"]'),
            ("柔韧性", '["拉伸", "瑜伽", "柔韧性", "活动度", "泡沫轴", "放松", "热身", "冷身"]'),
            ("损伤康复", '["受伤", "损伤", "康复", "疼痛", "扭伤", "拉伤", "骨折", "恢复", "理疗"]'),
            ("营养饮食", '["营养", "饮食", "蛋白质", "碳水", "脂肪", "补剂", "维生素", "矿物质", "食谱", "增肌餐", "减脂餐"]'),
            ("运动计划", '["训练计划", "分化训练", "新手", "进阶", "周期", "组数", "次数", "间歇", "容量", "强度"]'),
        ]
        cursor.executemany(
            "INSERT INTO category_keywords (category, keywords) VALUES (?, ?)",
            defaults,
        )

    conn.commit()
    conn.close()


def save_history(question: str, answer: str, mode: str, sources: str = None, user_id: int = None):
    """保存问答历史。匿名用户(user_id=None)不保存。"""
    if user_id is None:
        return
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO qa_history (question, answer, mode, sources, user_id) VALUES (?, ?, ?, ?, ?)",
        (question, answer, mode, sources, user_id)
    )
    conn.commit()
    conn.close()


def get_history(user_id: int, limit: int = 50):
    """获取当前用户的问答历史。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, question, answer, mode, sources, created_at FROM qa_history WHERE user_id = ? ORDER BY id DESC LIMIT ?",
        (user_id, limit)
    )
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            "id": row["id"],
            "question": row["question"],
            "answer": row["answer"],
            "mode": row["mode"],
            "sources": row["sources"],
            "created_at": row["created_at"],
        }
        for row in rows
    ]


def clear_history(user_id: int):
    """清空当前用户的问答历史。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM qa_history WHERE user_id = ?", (user_id,))
    conn.commit()
    conn.close()


def delete_history(history_id: int, user_id: int) -> bool:
    """删除当前用户的单条问答历史。成功返回 True。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM qa_history WHERE id = ? AND user_id = ?", (history_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    cursor.execute("DELETE FROM qa_history WHERE id = ? AND user_id = ?", (history_id, user_id))
    conn.commit()
    conn.close()
    return True


def delete_history_batch(ids: List[int], user_id: int) -> int:
    """批量删除当前用户的问答历史，返回实际删除条数。"""
    if not ids:
        return 0
    conn = get_connection()
    cursor = conn.cursor()
    placeholders = ",".join("?" * len(ids))
    cursor.execute(f"DELETE FROM qa_history WHERE id IN ({placeholders}) AND user_id = ?", ids + [user_id])
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted


# ==================== LLM 模型管理（按用户隔离，user_id=None 为全局） ====================

def _model_scope(user_id: Optional[int]):
    """返回模型表按用户过滤的 SQL 片段与参数。None 表示全局模型。"""
    if user_id is None:
        return "user_id IS NULL", ()
    return "user_id = ?", (user_id,)


def get_all_models(user_id: Optional[int] = None) -> List[Dict]:
    """获取指定范围的 LLM 模型列表（API Key 脱敏）。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(
        f"SELECT id, name, base_url, api_key_encrypted, model_name, is_active FROM llm_models WHERE {where} ORDER BY id",
        params,
    )
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            "id": row["id"],
            "name": row["name"],
            "base_url": row["base_url"],
            "api_key_masked": _mask_key(row["api_key_encrypted"]),
            "model_name": row["model_name"],
            "is_active": bool(row["is_active"]),
        }
        for row in rows
    ]


def get_active_model(user_id: Optional[int] = None) -> Optional[Dict]:
    """获取指定范围当前激活的模型（含加密 key）。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(
        f"SELECT id, name, base_url, api_key_encrypted, model_name, is_active FROM llm_models WHERE is_active = 1 AND {where} LIMIT 1",
        params,
    )
    row = cursor.fetchone()
    conn.close()
    if row:
        return {
            "id": row["id"],
            "name": row["name"],
            "base_url": row["base_url"],
            "api_key_encrypted": row["api_key_encrypted"],
            "model_name": row["model_name"],
            "is_active": bool(row["is_active"]),
        }
    return None


def get_model_by_id(model_id: int, user_id: Optional[int] = None) -> Optional[Dict]:
    """根据 ID 获取指定范围模型（含加密 key）。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(
        f"SELECT id, name, base_url, api_key_encrypted, model_name, is_active FROM llm_models WHERE id = ? AND {where}",
        (model_id,) + params,
    )
    row = cursor.fetchone()
    conn.close()
    if row:
        return {
            "id": row["id"],
            "name": row["name"],
            "base_url": row["base_url"],
            "api_key_encrypted": row["api_key_encrypted"],
            "model_name": row["model_name"],
            "is_active": bool(row["is_active"]),
        }
    return None


def add_model(name: str, base_url: str, api_key_encrypted: str, model_name: str, user_id: Optional[int] = None) -> int:
    """添加新模型，自动设为激活。返回新模型 ID。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(f"UPDATE llm_models SET is_active = 0 WHERE {where}", params)
    cursor.execute(
        "INSERT INTO llm_models (name, base_url, api_key_encrypted, model_name, is_active, user_id) VALUES (?, ?, ?, ?, 1, ?)",
        (name, base_url, api_key_encrypted, model_name, user_id)
    )
    model_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return model_id


def set_active_model(model_id: int, user_id: Optional[int] = None) -> bool:
    """切换指定范围的激活模型。成功返回 True。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(f"SELECT id FROM llm_models WHERE id = ? AND {where}", (model_id,) + params)
    if not cursor.fetchone():
        conn.close()
        return False
    cursor.execute(f"UPDATE llm_models SET is_active = 0 WHERE {where}", params)
    cursor.execute("UPDATE llm_models SET is_active = 1 WHERE id = ?", (model_id,))
    conn.commit()
    conn.close()
    return True


def delete_model(model_id: int, user_id: Optional[int] = None) -> bool:
    """删除指定范围模型。成功返回 True。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(f"SELECT id FROM llm_models WHERE id = ? AND {where}", (model_id,) + params)
    if not cursor.fetchone():
        conn.close()
        return False
    cursor.execute(f"DELETE FROM llm_models WHERE id = ? AND {where}", (model_id,) + params)
    conn.commit()
    conn.close()
    return True


def update_model(model_id: int, name: str, base_url: str, api_key_encrypted: str, model_name: str, user_id: Optional[int] = None) -> bool:
    """更新指定范围模型信息。成功返回 True。"""
    conn = get_connection()
    cursor = conn.cursor()
    where, params = _model_scope(user_id)
    cursor.execute(f"SELECT id FROM llm_models WHERE id = ? AND {where}", (model_id,) + params)
    if not cursor.fetchone():
        conn.close()
        return False
    cursor.execute(
        f"UPDATE llm_models SET name = ?, base_url = ?, api_key_encrypted = ?, model_name = ? WHERE id = ? AND {where}",
        (name, base_url, api_key_encrypted, model_name, model_id) + params
    )
    conn.commit()
    conn.close()
    return True


def _mask_key(encrypted_key: str) -> str:
    """对加密后的 key 做脱敏处理（解密后脱敏）。"""
    try:
        decrypted = settings.decrypt_api_key(encrypted_key)
        if len(decrypted) <= 8:
            return "****"
        return f"****{decrypted[-4:]}"
    except Exception:
        return "****"


# ==================== 动态知识条目管理 ====================

def add_knowledge_entries(entries: List[Dict], user_id: int = None) -> int:
    """批量插入知识条目，跳过已存在的 entry_id。返回新增数量。"""
    conn = get_connection()
    cursor = conn.cursor()
    added = 0
    for entry in entries:
        try:
            cursor.execute(
                """INSERT INTO knowledge_entries (entry_id, title, category, content, source, url, tags, import_mode, user_id, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    entry["entry_id"],
                    entry["title"],
                    entry.get("category", "未分类"),
                    entry["content"],
                    entry.get("source", ""),
                    entry.get("url", ""),
                    entry.get("tags", "[]"),
                    entry.get("import_mode", "direct"),
                    user_id,
                    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                )
            )
            added += 1
        except sqlite3.IntegrityError:
            continue
    conn.commit()
    conn.close()
    return added


def get_dynamic_entries(user_id: int = None) -> List[Dict]:
    """获取动态导入的知识条目。user_id=None 返回所有（用于检索），否则只返回当前用户的。"""
    import json
    conn = get_connection()
    cursor = conn.cursor()
    if user_id is not None:
        cursor.execute("SELECT entry_id, title, category, content, source, url, tags, created_at FROM knowledge_entries WHERE user_id = ? ORDER BY id", (user_id,))
    else:
        cursor.execute("SELECT entry_id, title, category, content, source, url, tags, created_at FROM knowledge_entries ORDER BY id")
    rows = cursor.fetchall()
    conn.close()
    result = []
    for row in rows:
        tags_raw = row["tags"] or "[]"
        try:
            tags = json.loads(tags_raw) if isinstance(tags_raw, str) else tags_raw
        except Exception:
            tags = []
        result.append({
            "id": row["entry_id"],
            "domain": row["category"],
            "title": row["title"],
            "category": row["category"],
            "content": row["content"],
            "source": row["source"],
            "url": row["url"] or "",
            "tags": tags,
            "chunk_id": row["entry_id"],
            "created_at": row["created_at"] or "",
        })
    return result


def get_dynamic_entries_list(user_id: int) -> List[Dict]:
    """获取当前用户的动态条目列表（含元数据）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT entry_id, title, category, content, source, import_mode, created_at FROM knowledge_entries WHERE user_id = ? ORDER BY id DESC", (user_id,))
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            "entry_id": row["entry_id"],
            "title": row["title"],
            "category": row["category"],
            "content": row["content"][:200] + ("..." if len(row["content"]) > 200 else ""),
            "source": row["source"],
            "import_mode": row["import_mode"],
            "created_at": row["created_at"],
        }
        for row in rows
    ]


def delete_entry(entry_id: str, user_id: int) -> bool:
    """删除当前用户的单条动态知识（同时删除版本记录）。成功返回 True。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM knowledge_entries WHERE entry_id = ? AND user_id = ?", (entry_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    cursor.execute("DELETE FROM knowledge_entries WHERE entry_id = ? AND user_id = ?", (entry_id, user_id))
    cursor.execute("DELETE FROM knowledge_versions WHERE entry_id = ?", (entry_id,))
    conn.commit()
    conn.close()
    return True


def is_source_imported(source: str, user_id: int) -> bool:
    """检查该来源文件是否已被当前用户导入。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM knowledge_entries WHERE source = ? AND user_id = ?", (source, user_id))
    count = cursor.fetchone()[0]
    conn.close()
    return count > 0


def add_single_entry(title: str, category: str, content: str, source: str = "", url: str = "", tags: str = "[]", user_id: int = None) -> str:
    """手动新增单条知识条目，返回新 entry_id。"""
    import json
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT MAX(CAST(SUBSTR(entry_id, 3) AS INTEGER)) FROM knowledge_entries WHERE entry_id LIKE 'E%'")
    row = cursor.fetchone()
    next_num = (row[0] or 0) + 1
    entry_id = f"E{next_num:03d}"

    cursor.execute(
        """INSERT INTO knowledge_entries (entry_id, title, category, content, source, url, tags, import_mode, user_id, created_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (entry_id, title, category, content, source, url, tags, "manual", user_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    conn.commit()
    conn.close()
    return entry_id


def update_entry(entry_id: str, title: str, category: str, content: str, source: str = "", url: str = "", tags: str = "[]", user_id: int = None) -> bool:
    """编辑当前用户的知识条目（自动保存旧版本）。成功返回 True。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, title, category, content, source, url FROM knowledge_entries WHERE entry_id = ? AND user_id = ?", (entry_id, user_id))
    old = cursor.fetchone()
    if not old:
        conn.close()
        return False

    # 保存旧版本
    save_knowledge_version(
        entry_id=entry_id,
        title=old["title"],
        category=old["category"],
        content=old["content"],
        source=old["source"] or "",
        url=old["url"] or "",
    )

    cursor.execute(
        """UPDATE knowledge_entries SET title=?, category=?, content=?, source=?, url=?, tags=? WHERE entry_id=? AND user_id=?""",
        (title, category, content, source, url, tags, entry_id, user_id)
    )
    conn.commit()
    conn.close()
    return True


def get_entry_by_id(entry_id: str, user_id: int = None) -> Optional[Dict]:
    """根据 entry_id 获取单条知识。user_id 不为 None 时校验归属。"""
    conn = get_connection()
    cursor = conn.cursor()
    if user_id is not None:
        cursor.execute("SELECT * FROM knowledge_entries WHERE entry_id = ? AND user_id = ?", (entry_id, user_id))
    else:
        cursor.execute("SELECT * FROM knowledge_entries WHERE entry_id = ?", (entry_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    return dict(row)


def update_entries_batch(entry_ids: List[str], fields: Dict) -> int:
    """
    批量更新知识条目的公共字段（仅动态条目）。

    fields 支持: category / source / url / tags（tags 为 JSON 字符串）。
    只更新 fields 中出现的字段，返回实际更新条数。
    """
    if not entry_ids or not fields:
        return 0

    allowed = {"category", "source", "url", "tags"}
    updates = {k: v for k, v in fields.items() if k in allowed}
    if not updates:
        return 0

    set_clause = ", ".join(f"{k}=?" for k in updates)
    placeholders = ",".join("?" * len(entry_ids))
    sql = f"UPDATE knowledge_entries SET {set_clause} WHERE entry_id IN ({placeholders})"

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(sql, list(updates.values()) + entry_ids)
    updated = cursor.rowcount
    conn.commit()
    conn.close()
    return updated


def find_similar_entries(title: str, content: str, threshold: float = 0.85, user_id: int = None) -> List[Dict]:
    """
    查找与给定标题/内容高度相似（重复）的已有条目。
    user_id 不为 None 时只查当前用户的条目。

    使用 difflib 计算标题与内容文本的相似度，超过阈值视为重复。
    返回相似条目的列表（含 entry_id、title、similarity）。
    """
    import difflib
    conn = get_connection()
    cursor = conn.cursor()
    if user_id is not None:
        cursor.execute("SELECT entry_id, title, content FROM knowledge_entries WHERE user_id = ?", (user_id,))
    else:
        cursor.execute("SELECT entry_id, title, content FROM knowledge_entries")
    rows = cursor.fetchall()
    conn.close()

    new_text = f"{title}\n{content}"
    similar = []
    for row in rows:
        existing_text = f"{row['title']}\n{row['content']}"
        # 标题精确匹配视为重复
        if row["title"] and title and row["title"].strip() == title.strip():
            similar.append({
                "entry_id": row["entry_id"],
                "title": row["title"],
                "similarity": 1.0,
            })
            continue
        ratio = difflib.SequenceMatcher(None, new_text, existing_text).ratio()
        if ratio >= threshold:
            similar.append({
                "entry_id": row["entry_id"],
                "title": row["title"],
                "similarity": round(ratio, 4),
            })
    return similar


# ==================== 用户认证 ====================

def create_user(username: str, password_hash: str) -> Optional[int]:
    """创建用户。用户名重复返回 None。"""
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO users (username, password_hash) VALUES (?, ?)",
            (username, password_hash),
        )
        uid = cursor.lastrowid
        conn.commit()
        return uid
    except sqlite3.IntegrityError:
        return None
    finally:
        conn.close()


def get_user_by_username(username: str) -> Optional[Dict]:
    """根据用户名获取用户。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, password_hash, created_at FROM users WHERE username = ?", (username,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    return dict(row)


def create_auth_token(user_id: int) -> str:
    """为用户生成登录令牌（7天过期）。"""
    import secrets
    from datetime import timedelta
    token = secrets.token_hex(32)
    expires_at = (datetime.now() + timedelta(days=settings.TOKEN_EXPIRE_DAYS)).strftime("%Y-%m-%d %H:%M:%S")
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO auth_tokens (token, user_id, expires_at) VALUES (?, ?, ?)",
        (token, user_id, expires_at),
    )
    conn.commit()
    conn.close()
    return token


def get_user_by_token(token: str) -> Optional[Dict]:
    """根据令牌获取用户（过期或无效返回 None）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT u.id, u.username, t.expires_at FROM auth_tokens t JOIN users u ON u.id = t.user_id WHERE t.token = ?",
        (token,),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    if row["expires_at"]:
        from datetime import datetime as _dt
        if _dt.strptime(row["expires_at"], "%Y-%m-%d %H:%M:%S") < _dt.now():
            return None
    return {"id": row["id"], "username": row["username"]}


def cleanup_expired_tokens() -> int:
    """清理过期的 auth_tokens。返回删除数量。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "DELETE FROM auth_tokens WHERE expires_at IS NOT NULL AND expires_at < datetime('now', 'localtime')"
    )
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted


def delete_auth_token(token: str) -> bool:
    """删除登录令牌（退出登录）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM auth_tokens WHERE token = ?", (token,))
    deleted = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return deleted


# ==================== 多轮会话 ====================

def create_chat_session(user_id: int, title: str = "新对话") -> int:
    """创建会话，返回会话 ID。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO chat_sessions (user_id, title) VALUES (?, ?)",
        (user_id, title),
    )
    sid = cursor.lastrowid
    conn.commit()
    conn.close()
    return sid


def list_chat_sessions(user_id: int) -> List[Dict]:
    """获取用户的所有会话（按更新时间倒序）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT id, title, created_at, updated_at,
                  (SELECT COUNT(*) FROM chat_messages m WHERE m.session_id = chat_sessions.id) AS message_count
           FROM chat_sessions WHERE user_id = ? ORDER BY updated_at DESC""",
        (user_id,),
    )
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            "id": row["id"],
            "title": row["title"],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
            "message_count": row["message_count"],
        }
        for row in rows
    ]


def get_chat_session(user_id: int, session_id: int) -> Optional[Dict]:
    """获取属于该用户的会话。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, title, created_at, updated_at FROM chat_sessions WHERE id = ? AND user_id = ?", (session_id, user_id))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    return dict(row)


def update_chat_session_title(session_id: int, title: str) -> bool:
    """更新会话标题（仅在需要时调用）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE chat_sessions SET title = ? WHERE id = ?", (title[:100], session_id))
    changed = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return changed


def delete_chat_session(user_id: int, session_id: int) -> bool:
    """删除会话（同时删除消息）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM chat_sessions WHERE id = ? AND user_id = ?", (session_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    cursor.execute("DELETE FROM chat_messages WHERE session_id = ?", (session_id,))
    cursor.execute("DELETE FROM chat_sessions WHERE id = ?", (session_id,))
    conn.commit()
    conn.close()
    return True


def add_chat_message(session_id: int, role: str, content: str, sources: str = None) -> int:
    """追加一条会话消息，并刷新会话更新时间。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO chat_messages (session_id, role, content, sources) VALUES (?, ?, ?, ?)",
        (session_id, role, content, sources),
    )
    mid = cursor.lastrowid
    cursor.execute("UPDATE chat_sessions SET updated_at = CURRENT_TIMESTAMP WHERE id = ?", (session_id,))
    conn.commit()
    conn.close()
    return mid


def list_chat_messages(session_id: int) -> List[Dict]:
    """获取会话的全部消息。"""
    import json as _json
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, role, content, sources, created_at FROM chat_messages WHERE session_id = ? ORDER BY id",
        (session_id,),
    )
    rows = cursor.fetchall()
    conn.close()
    messages = []
    for row in rows:
        try:
            sources = _json.loads(row["sources"]) if row["sources"] else []
        except Exception:
            sources = []
        messages.append({
            "id": row["id"],
            "role": row["role"],
            "content": row["content"],
            "sources": sources,
            "created_at": row["created_at"],
        })
    return messages


# ==================== 无法回答问题记录 ====================

def record_unanswered_question(question: str, mode: str = "hybrid", reason: str = ""):
    """记录无法回答的问题（无检索结果或分数过低）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO unanswered_questions (question, mode, reason) VALUES (?, ?, ?)",
        (question, mode, reason),
    )
    conn.commit()
    conn.close()


def list_unanswered_questions(limit: int = 100) -> List[Dict]:
    """获取无法回答的问题记录。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, question, mode, reason, created_at FROM unanswered_questions ORDER BY id DESC LIMIT ?",
        (limit,),
    )
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


def clear_unanswered_questions() -> int:
    """清空无法回答问题记录。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM unanswered_questions")
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted


# ==================== 每用户 LLM 配置 ====================

def get_user_llm_config(user_id: int) -> Optional[Dict]:
    """获取指定用户的 LLM 配置（无则返回 None）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT base_url, api_key_encrypted, model_name, mode FROM user_llm_config WHERE user_id = ?",
        (user_id,),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    return dict(row)


def save_user_llm_config(user_id: int, base_url: str, api_key_encrypted: str, model_name: str, mode: str):
    """保存/更新用户的 LLM 配置。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO user_llm_config (user_id, base_url, api_key_encrypted, model_name, mode, updated_at)
        VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(user_id) DO UPDATE SET
            base_url=excluded.base_url,
            api_key_encrypted=excluded.api_key_encrypted,
            model_name=excluded.model_name,
            mode=excluded.mode,
            updated_at=CURRENT_TIMESTAMP
        """,
        (user_id, base_url, api_key_encrypted, model_name, mode),
    )
    conn.commit()
    conn.close()


def get_user_mode(user_id: int) -> str:
    """获取用户的运行模式（默认 mock）。"""
    cfg = get_user_llm_config(user_id)
    mode = (cfg or {}).get("mode")
    return mode if mode in ("mock", "real") else "mock"


def set_user_mode(user_id: int, mode: str):
    """设置用户的运行模式。"""
    if mode not in ("mock", "real"):
        mode = "mock"
    save_user_llm_config(user_id, "", "", "", mode)


# ==================== 运动记录 ====================

def add_exercise_record(user_id: int, exercise_type: str, duration: int = 0,
                        intensity: str = "中等", notes: str = "", record_date: str = "") -> int:
    """添加运动记录，返回记录 ID。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """INSERT INTO exercise_records (user_id, exercise_type, duration, intensity, notes, record_date)
           VALUES (?, ?, ?, ?, ?, ?)""",
        (user_id, exercise_type, duration, intensity, notes, record_date),
    )
    rid = cursor.lastrowid
    conn.commit()
    conn.close()
    return rid


def list_exercise_records(user_id: int, limit: int = 100, offset: int = 0,
                          exercise_type: str = "", duration_min: int = 0, duration_max: int = 0,
                          intensity: str = "", date_from: str = "", date_to: str = "",
                          keyword: str = "") -> List[Dict]:
    """获取用户运动记录，支持筛选。"""
    conn = get_connection()
    cursor = conn.cursor()
    conditions = ["user_id = ?"]
    params = [user_id]
    if exercise_type:
        conditions.append("exercise_type LIKE ?")
        params.append(f"%{exercise_type}%")
    if intensity:
        conditions.append("intensity = ?")
        params.append(intensity)
    if duration_min > 0:
        conditions.append("duration >= ?")
        params.append(duration_min)
    if duration_max > 0:
        conditions.append("duration <= ?")
        params.append(duration_max)
    if date_from:
        conditions.append("date(record_date) >= ?")
        params.append(date_from)
    if date_to:
        conditions.append("date(record_date) <= ?")
        params.append(date_to)
    if keyword:
        conditions.append("(exercise_type LIKE ? OR notes LIKE ?)")
        kw = f"%{keyword}%"
        params.extend([kw, kw])
    where = " AND ".join(conditions)
    cursor.execute(
        f"""SELECT id, exercise_type, duration, intensity, notes, record_date, created_at
           FROM exercise_records
           WHERE {where}
           ORDER BY record_date DESC, id DESC
           LIMIT ? OFFSET ?""",
        (*params, limit, offset),
    )
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


def update_exercise_record(user_id: int, record_id: int, exercise_type: str = "",
                           duration: int = 0, intensity: str = "", notes: str = "",
                           record_date: str = "") -> bool:
    """更新运动记录。"""
    conn = get_connection()
    cursor = conn.cursor()
    existing = cursor.execute(
        "SELECT * FROM exercise_records WHERE id = ? AND user_id = ?",
        (record_id, user_id),
    ).fetchone()
    if not existing:
        conn.close()
        return False
    new_type = exercise_type if exercise_type else existing["exercise_type"]
    new_duration = duration if duration else existing["duration"]
    new_intensity = intensity if intensity else existing["intensity"]
    new_notes = notes if notes else existing["notes"]
    new_date = record_date if record_date else existing["record_date"]
    cursor.execute(
        """UPDATE exercise_records SET exercise_type=?, duration=?, intensity=?, notes=?, record_date=?
           WHERE id=? AND user_id=?""",
        (new_type, new_duration, new_intensity, new_notes, new_date, record_id, user_id),
    )
    updated = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return updated


def delete_exercise_record(user_id: int, record_id: int) -> bool:
    """删除指定用户的运动记录。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "DELETE FROM exercise_records WHERE id = ? AND user_id = ?",
        (record_id, user_id),
    )
    deleted = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return deleted


def batch_delete_exercise_records(record_ids: list, user_id: int) -> int:
    """批量删除运动记录，返回删除数量。"""
    conn = get_connection()
    cursor = conn.cursor()
    placeholders = ",".join("?" for _ in record_ids)
    cursor.execute(
        f"DELETE FROM exercise_records WHERE id IN ({placeholders}) AND user_id = ?",
        (*record_ids, user_id),
    )
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted


# ==================== 训练计划 ====================

def add_training_plan(user_id: int, title: str, goal: str, level: str,
                      days_per_week: int, content: str, category: str = "") -> int:
    """添加训练计划，返回 ID。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """INSERT INTO training_plans (user_id, title, goal, level, days_per_week, content, category)
           VALUES (?, ?, ?, ?, ?, ?, ?)""",
        (user_id, title, goal, level, days_per_week, content, category),
    )
    rid = cursor.lastrowid
    conn.commit()
    conn.close()
    return rid


def update_training_plan(plan_id: int, user_id: int, title: str, goal: str, level: str,
                         days_per_week: int, content: str, category: str = "") -> bool:
    """更新训练计划。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """UPDATE training_plans SET title=?, goal=?, level=?, days_per_week=?, content=?, category=?,
           updated_at=CURRENT_TIMESTAMP WHERE id=? AND user_id=?""",
        (title, goal, level, days_per_week, content, category, plan_id, user_id),
    )
    updated = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return updated


def delete_training_plan(plan_id: int, user_id: int) -> bool:
    """删除单个训练计划。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "DELETE FROM training_plans WHERE id = ? AND user_id = ?",
        (plan_id, user_id),
    )
    deleted = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return deleted


def batch_delete_training_plans(plan_ids: list, user_id: int) -> int:
    """批量删除训练计划，返回删除数量。"""
    conn = get_connection()
    cursor = conn.cursor()
    placeholders = ",".join("?" for _ in plan_ids)
    cursor.execute(
        f"DELETE FROM training_plans WHERE id IN ({placeholders}) AND user_id = ?",
        (*plan_ids, user_id),
    )
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted


def batch_update_training_plans(plan_ids: list, user_id: int, goal: str = "",
                                level: str = "", days_per_week: int = 0,
                                category: str = "") -> int:
    """批量更新训练计划，返回更新数量。"""
    conn = get_connection()
    cursor = conn.cursor()
    set_parts = []
    params = []
    if goal:
        set_parts.append("goal = ?")
        params.append(goal)
    if level:
        set_parts.append("level = ?")
        params.append(level)
    if days_per_week > 0:
        set_parts.append("days_per_week = ?")
        params.append(days_per_week)
    if category:
        set_parts.append("category = ?")
        params.append(category)
    if not set_parts:
        conn.close()
        return 0
    set_parts.append("updated_at = CURRENT_TIMESTAMP")
    placeholders = ",".join("?" for _ in plan_ids)
    params.extend([*plan_ids, user_id])
    cursor.execute(
        f"UPDATE training_plans SET {', '.join(set_parts)} "
        f"WHERE id IN ({placeholders}) AND user_id = ?",
        params,
    )
    updated = cursor.rowcount
    conn.commit()
    conn.close()
    return updated


def list_training_plans(user_id: int, category: str = "", date_from: str = "", date_to: str = "",
                        keyword: str = "", goal: str = "", level: str = "",
                        days_per_week: int = 0, limit: int = 100, offset: int = 0) -> List[Dict]:
    """列出训练计划，支持筛选。"""
    conn = get_connection()
    cursor = conn.cursor()
    conditions = ["user_id = ?"]
    params = [user_id]
    if category:
        conditions.append("category = ?")
        params.append(category)
    if goal:
        conditions.append("goal LIKE ?")
        params.append(f"%{goal}%")
    if level:
        conditions.append("level = ?")
        params.append(level)
    if days_per_week > 0:
        conditions.append("days_per_week = ?")
        params.append(days_per_week)
    if date_from:
        conditions.append("date(created_at) >= ?")
        params.append(date_from)
    if date_to:
        conditions.append("date(created_at) <= ?")
        params.append(date_to)
    if keyword:
        conditions.append("(title LIKE ? OR content LIKE ? OR goal LIKE ?)")
        kw = f"%{keyword}%"
        params.extend([kw, kw, kw])
    where = " AND ".join(conditions)
    cursor.execute(
        f"SELECT id, title, goal, level, days_per_week, content, category, created_at, updated_at "
        f"FROM training_plans WHERE {where} ORDER BY updated_at DESC, created_at DESC LIMIT ? OFFSET ?",
        (*params, limit, offset),
    )
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


def get_training_plan(plan_id: int, user_id: int) -> Optional[Dict]:
    """获取单个训练计划详情。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, title, goal, level, days_per_week, content, category, created_at, updated_at "
        "FROM training_plans WHERE id = ? AND user_id = ?",
        (plan_id, user_id),
    )
    row = cursor.fetchone()
    conn.close()
    return dict(row) if row else None


def get_training_plan_categories(user_id: int) -> List[str]:
    """获取用户训练计划的所有分类。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT DISTINCT category FROM training_plans WHERE user_id = ? AND category != '' ORDER BY category",
        (user_id,),
    )
    rows = cursor.fetchall()
    conn.close()
    return [row["category"] for row in rows]


# ==================== 知识版本管理 ====================

def save_knowledge_version(entry_id: str, title: str, category: str, content: str, source: str = "", url: str = "") -> int:
    """保存知识条目版本。返回新版本号。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT MAX(version) FROM knowledge_versions WHERE entry_id = ?",
        (entry_id,),
    )
    row = cursor.fetchone()
    new_version = (row[0] or 0) + 1
    cursor.execute(
        """INSERT INTO knowledge_versions (entry_id, version, title, category, content, source, url)
           VALUES (?, ?, ?, ?, ?, ?, ?)""",
        (entry_id, new_version, title, category, content, source, url),
    )
    conn.commit()
    conn.close()
    return new_version


def get_knowledge_versions(entry_id: str) -> List[Dict]:
    """获取知识条目的所有版本（按版本号倒序）。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT id, entry_id, version, title, category, content, source, url, created_at
           FROM knowledge_versions WHERE entry_id = ? ORDER BY version DESC""",
        (entry_id,),
    )
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


def get_knowledge_version(version_id: int) -> Optional[Dict]:
    """获取指定版本详情。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, entry_id, version, title, category, content, source, url, created_at FROM knowledge_versions WHERE id = ?",
        (version_id,),
    )
    row = cursor.fetchone()
    conn.close()
    return dict(row) if row else None


def restore_knowledge_version(version_id: int) -> bool:
    """将知识条目恢复到指定版本。"""
    version = get_knowledge_version(version_id)
    if not version:
        return False
    return update_entry(
        version["entry_id"],
        version["title"],
        version["category"],
        version["content"],
        version["source"],
        version["url"] or "",
    )


def delete_knowledge_versions(entry_id: str) -> int:
    """删除知识条目的所有版本记录。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM knowledge_versions WHERE entry_id = ?", (entry_id,))
    deleted = cursor.rowcount
    conn.commit()
    conn.close()
    return deleted


# ==================== 检索偏好 ====================

def get_user_retrieval_config(user_id: int) -> Optional[Dict]:
    """获取用户检索偏好。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT default_mode, default_top_k, min_vector_score FROM user_retrieval_config WHERE user_id = ?",
        (user_id,),
    )
    row = cursor.fetchone()
    conn.close()
    return dict(row) if row else None


def save_user_retrieval_config(user_id: int, default_mode: str, default_top_k: int, min_vector_score: float):
    """保存/更新用户检索偏好。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO user_retrieval_config (user_id, default_mode, default_top_k, min_vector_score, updated_at)
        VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(user_id) DO UPDATE SET
            default_mode=excluded.default_mode,
            default_top_k=excluded.default_top_k,
            min_vector_score=excluded.min_vector_score,
            updated_at=CURRENT_TIMESTAMP
        """,
        (user_id, default_mode, default_top_k, min_vector_score),
    )
    conn.commit()
    conn.close()


# ==================== 分类关键词管理 ====================

def get_all_category_keywords() -> List[Dict]:
    """获取所有分类及其关键词。"""
    import json as _json
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT category, keywords FROM category_keywords ORDER BY category")
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            "category": row["category"],
            "keywords": _json.loads(row["keywords"]) if row["keywords"] else [],
        }
        for row in rows
    ]


def save_category_keywords(category: str, keywords: List[str]):
    """保存/更新分类关键词。"""
    import json as _json
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO category_keywords (category, keywords, updated_at)
        VALUES (?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(category) DO UPDATE SET
            keywords=excluded.keywords,
            updated_at=CURRENT_TIMESTAMP
        """,
        (category, _json.dumps(keywords, ensure_ascii=False)),
    )
    conn.commit()
    conn.close()


def delete_category(category: str) -> bool:
    """删除分类。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM category_keywords WHERE category = ?", (category,))
    deleted = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return deleted


def get_category_names() -> List[str]:
    """获取所有分类名称列表。"""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT category FROM category_keywords ORDER BY category")
    rows = cursor.fetchall()
    conn.close()
    return [row["category"] for row in rows]


# ==================== 知识导出 ====================

def export_knowledge(format: str = "json") -> str:
    """导出知识库为 JSON / CSV / Markdown / TXT 字符串（PDF 请使用 export_knowledge_pdf）。"""
    import json as _json
    import csv
    import io

    entries = get_dynamic_entries()
    if format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["ID", "标题", "分类", "内容", "来源", "URL", "标签", "创建时间"])
        for e in entries:
            writer.writerow([
                e["id"], e["title"], e["category"], e["content"],
                e["source"], e["url"], ", ".join(e.get("tags", [])),
                e["created_at"],
            ])
        return output.getvalue()
    elif format == "markdown":
        lines = [f"# FitQA 知识库导出\n\n> 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n---\n"]
        for e in entries:
            lines.append(f"## {e['title']}\n\n")
            lines.append(f"- **分类**：{e['category']}\n")
            lines.append(f"- **来源**：{e['source']}\n")
            if e["url"]:
                lines.append(f"- **链接**：{e['url']}\n")
            lines.append(f"\n{e['content']}\n\n---\n")
        return "".join(lines)
    elif format == "txt":
        lines = [f"FitQA 知识库导出\n", f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n", "=" * 50 + "\n\n"]
        for e in entries:
            lines.append(f"标题：{e['title']}\n")
            lines.append(f"分类：{e['category']}\n")
            if e["source"]:
                lines.append(f"来源：{e['source']}\n")
            if e["url"]:
                lines.append(f"链接：{e['url']}\n")
            lines.append(f"\n{e['content']}\n")
            lines.append("-" * 50 + "\n\n")
        return "".join(lines)
    else:
        return _json.dumps(entries, ensure_ascii=False, indent=2)


def export_knowledge_docx() -> bytes:
    """导出知识库为 Word (.docx) 字节流。"""
    from docx import Document

    doc = Document()
    doc.add_heading("FitQA 知识库导出", level=0)
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("—" * 40)

    entries = get_dynamic_entries()
    for e in entries:
        doc.add_heading(e["title"], level=2)
        doc.add_paragraph(f"分类：{e['category']}")
        doc.add_paragraph(f"来源：{e['source']}")
        if e["url"]:
            doc.add_paragraph(f"链接：{e['url']}")
        doc.add_paragraph(e["content"])
        doc.add_paragraph("—" * 40)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_knowledge_pdf() -> bytes:
    """导出知识库为 PDF 字节流。"""
    from fpdf import FPDF
    import io

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Use built-in core font (Helvetica) for maximum compatibility
    # Chinese characters will render as-is (UTF-8) in fpdf2 with core fonts
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "FitQA Knowledge Base Export", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 8)
    from datetime import datetime
    pdf.cell(0, 6, f"Export Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    entries = get_dynamic_entries()
    for e in entries:
        # Check if we need a new page
        if pdf.get_y() > 250:
            pdf.add_page()
        pdf.set_font("Helvetica", "B", 12)
        # Write title
        pdf.multi_cell(0, 7, e["title"])
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, f"Category: {e['category']}", new_x="LMARGIN", new_y="NEXT")
        if e["source"]:
            pdf.cell(0, 5, f"Source: {e['source']}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        # Write content
        content = e["content"].replace("\n", " ").replace("\r", " ")
        pdf.multi_cell(0, 5, content)
        pdf.ln(3)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.getvalue()